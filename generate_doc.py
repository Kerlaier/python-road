# -*- coding: utf-8 -*-
"""Python 学习录 - 文档生成器"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os, re, shutil

SOURCE = os.path.expanduser(r"~\Desktop\python学习文件")
REPO = os.path.expanduser(r"~\Desktop\python-road")
SOURCE_DIR = os.path.join(REPO, "源代码")
OUTPUT = os.path.join(REPO, "学习录")

THEME_BLUE  = RGBColor(0x2B, 0x5C, 0x8A)
THEME_GOLD  = RGBColor(0xE8, 0x9C, 0x25)
THEME_GRAY  = RGBColor(0x3C, 0x3C, 0x3C)
THEME_LIGHT = RGBColor(0xF0, 0xF4, 0xF8)
THEME_RED   = RGBColor(0xC0, 0x39, 0x2B)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_code_block(doc, code_text, title=""):
    if title:
        p = doc.add_paragraph()
        run = p.add_run(f"📄 {title}")
        run.bold = True; run.font.size = Pt(11); run.font.color.rgb = THEME_BLUE
    for line in code_text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line)
        run.font.name = "Courier New"; run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F4F4"/>')
        pPr.append(shading)

def add_tip_card(doc, tip_type, text):
    icons = {"避坑": "⚠️", "进阶": "🔥", "笔记": "💡"}
    colors = {"避坑": "FDEDEC", "进阶": "FEF9E7", "笔记": "EBF5FB"}
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, colors.get(tip_type, "F5F5F5"))
    p = cell.paragraphs[0]
    run = p.add_run(f"{icons.get(tip_type, '')} {tip_type}：{text}")
    run.font.size = Pt(10); run.font.name = "Microsoft YaHei"
    run.font.color.rgb = THEME_GRAY
    doc.add_paragraph()

def add_chapter_banner(doc, number, title, date):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("💡 建议配合 Python 运行本课讲义代码以加深理解")
    run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x99, 0x99, 0x99); run.italic = True
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"第 {number} 课")
    run.font.size = Pt(14); run.font.color.rgb = THEME_GOLD; run.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(26); run.font.color.rgb = THEME_BLUE; run.bold = True
    run.font.name = "Microsoft YaHei"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 40)
    run.font.color.rgb = THEME_GOLD; run.font.size = Pt(8)

def add_section_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(16); run.font.color.rgb = THEME_BLUE
    run.font.name = "Microsoft YaHei"
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="4" w:color="2B5C8A"/></w:pBdr>')
    pPr.append(pBdr)

def add_knowledge_card(doc, formula, code, output, summary):
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        row.cells[0].width = Cm(3)
        row.cells[1].width = Cm(13)
    set_cell_shading(table.cell(0, 0), "2B5C8A")
    p = table.cell(0, 0).paragraphs[0]
    run = p.add_run("语法公式"); run.bold = True; run.font.size = Pt(10); run.font.color.rgb = WHITE
    p = table.cell(0, 1).paragraphs[0]
    run = p.add_run(formula); run.font.name = "Courier New"; run.font.size = Pt(10)
    set_cell_shading(table.cell(1, 0), "F0F4F8")
    p = table.cell(1, 0).paragraphs[0]
    run = p.add_run("代码示例"); run.bold = True; run.font.size = Pt(10); run.font.color.rgb = THEME_BLUE
    p = table.cell(1, 1).paragraphs[0]
    run = p.add_run(code); run.font.name = "Courier New"; run.font.size = Pt(9.5)
    set_cell_shading(table.cell(2, 0), "F0F4F8")
    p = table.cell(2, 0).paragraphs[0]
    run = p.add_run("输出结果"); run.bold = True; run.font.size = Pt(10); run.font.color.rgb = THEME_BLUE
    p = table.cell(2, 1).paragraphs[0]
    run = p.add_run(output); run.font.size = Pt(9.5)
    set_cell_shading(table.cell(3, 0), "E89C25")
    p = table.cell(3, 0).paragraphs[0]
    run = p.add_run("一句话总结"); run.bold = True; run.font.size = Pt(10); run.font.color.rgb = WHITE
    p = table.cell(3, 1).paragraphs[0]
    run = p.add_run(summary); run.font.size = Pt(10); run.font.color.rgb = THEME_GRAY
    doc.add_paragraph()

def add_homework_item(doc, num, question, code, note=""):
    """添加单个作业：题目 + 答案代码"""
    # 作业标题
    p = doc.add_paragraph()
    run = p.add_run(f"作业{num}：{question}")
    run.bold = True; run.font.size = Pt(11); run.font.color.rgb = THEME_BLUE
    run.font.name = "Microsoft YaHei"
    # 提示
    if note:
        p = doc.add_paragraph()
        run = p.add_run(f"💡 {note}")
        run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        run.italic = True
    # 答案代码
    for line in code.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(f"  {line}")
        run.font.name = "Courier New"; run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0FAF0"/>')
        pPr.append(shading)
    doc.add_paragraph()  # 间距

def add_checklist(doc, items):
    add_section_title(doc, "✅ 章末 Checklist")
    for item in items:
        p = doc.add_paragraph(f"  ☐ {item}")
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)


def parse_practice_answers(lesson_number):
    """从 practice.py 解析作业答案，支持两种格式"""
    fname = f"lesson_{lesson_number:02d}_practice.py"
    filepath = os.path.join(SOURCE, fname)
    if not os.path.exists(filepath):
        print(f"  [警告] practice 文件不存在: {fname}")
        return []
    with open(filepath, "r", encoding="utf-8") as f:
        plines = f.readlines()
    num_map = {1: "一", 2: "二", 3: "三", 4: "四", 5: "五"}
    is_block = any(re.match(r"# ={20,}", l.strip()) for l in plines)
    exercises = []
    i = 0
    while i < len(plines):
        s = plines[i].strip()
        if is_block:
            m = re.match(r"# 练习(\d+)", s)
        else:
            m = re.match(r"# ={5,} 练习(\d+)", s)
        if m:
            num = int(m.group(1))
            if is_block:
                title_m = re.search(r"练习\d+[：:]\s*(.+)", s)
                question = title_m.group(1) if title_m else f"练习{num}"
                # 跳过到闭合分隔符 # ===== 之后
                j = i + 1
                while j < len(plines):
                    if re.match(r"# ={20,}", plines[j].strip()):
                        j += 1
                        break
                    if plines[j].strip().startswith("# 提示") or plines[j].strip().startswith("#提示"):
                        pass
                    j += 1
            else:
                title_m = re.search(r"练习\d+[：:]\s*(.+?)\s*=*$", s)
                question = title_m.group(1) if title_m else f"练习{num}"
                j = i + 1
            hint = ""
            code_lines = []
            while j < len(plines):
                sj = plines[j].strip()
                if is_block:
                    if re.match(r"# ={20,}", sj) or re.match(r"# 练习\d+", sj):
                        break
                else:
                    if re.match(r"# ={5,} 练习\d+", sj):
                        break
                if sj.startswith("# 提示") or sj.startswith("#提示"):
                    hint = sj.lstrip("# ").strip()
                    j += 1
                elif sj.startswith("#") or sj == "":
                    j += 1
                else:
                    code_lines.append(plines[j].rstrip())
                    j += 1
            code = "\n".join(code_lines)
            if code.strip():
                exercises.append((num_map.get(num, str(num)), question, code, hint))
        i += 1
    if not exercises:
        print(f"  [警告] practice 文件中未找到答案代码: {fname}")
    return exercises
def add_watermark(doc, text="by Kerlaier"):
    """添加隐秘水印 -- 浅灰色斜体页脚文字"""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xD0, 0xD0, 0xD0)
    run.font.italic = True
    run.font.name = "Microsoft YaHei"


def copy_source_files(lesson_number):
    """复制对应课程的 .py 源文件到仓库源代码目录"""
    lesson_map = {
        2: ["lesson_02_datatypes.py", "lesson_02_practice.py"],
        3: ["lesson_03_io.py", "lesson_03_practice.py"],
        4: ["lesson_04_if.py", "lesson_04_practice.py"],
        5: ["lesson_05_loop.py", "lesson_05_practice.py"],
        6: ["lesson_06_list.py", "lesson_06_practice.py"],
    }
    files = lesson_map.get(lesson_number, [])
    dest_dir = os.path.join(SOURCE_DIR, f"第{lesson_number}课")
    os.makedirs(dest_dir, exist_ok=True)
    for fname in files:
        src = os.path.join(SOURCE, fname)
        dst = os.path.join(dest_dir, fname)
        if os.path.exists(src):
            shutil.copy2(src, dst)
            print(f"  已复制源文件: {fname}")
        else:
            print(f"  [警告] 源文件不存在: {fname}")


def generate_lesson_2():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21); section.page_height = Cm(29.7)
    section.left_margin = Cm(2); section.right_margin = Cm(2)
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2)

    add_chapter_banner(doc, 2, "数据类型", "2026-07-14")

    p = doc.add_paragraph()
    run = p.add_run("🎯 本章目标")
    run.bold = True; run.font.size = Pt(13); run.font.color.rgb = THEME_BLUE
    for g in ["掌握 Python 四种核心数据类型：整数、浮点数、字符串、布尔值",
              "学会类型转换，避免 \"5\"+3 这类常见坑",
              "用 f-string 优雅地格式化输出"]:
        p = doc.add_paragraph(f"  ✓ {g}")
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)

    # 一、数字类型
    add_section_title(doc, "一、数字类型：int 与 float")
    add_knowledge_card(doc,
        "int: 整数（age = 25）\nfloat: 小数（pi = 3.14）",
        "age = 25\ntemperature = -10\npi = 3.1415926\nprint(type(age))  # <class 'int'>",
        ">>> 年龄: 25, 类型: <class 'int'>\n>>> 圆周率: 3.1415926",
        "整数用 int，小数用 float，Python 自动区分类型")
    add_code_block(doc,
        "a, b = 10, 3\n"
        "print(f\"加法: {a + b}\")        # 13\n"
        "print(f\"除法: {a / b}\")        # 3.333... (永远返回 float)\n"
        "print(f\"整除: {a // b}\")       # 3 (去掉小数部分)\n"
        "print(f\"取余: {a % b}\")        # 1 (余数)\n"
        "print(f\"幂: {a ** b}\")         # 1000 (a的b次方)\n"
        "\n# 常用数学函数\n"
        "abs(-5)        → 5        # 绝对值\n"
        "round(3.14159, 2) → 3.14 # 保留2位小数\n"
        "max(1,5,3)     → 5        # 最大值\n"
        "pow(2, 10)     → 1024     # 2的10次方",
        "基本运算速查")
    add_tip_card(doc, "避坑", "除法 / 永远返回 float，即使能整除！7 / 2 = 3.5，不是 3")

    # 二、字符串
    add_section_title(doc, "二、字符串：str")
    add_knowledge_card(doc,
        "单引号 / 双引号 / 三引号多行\n索引 word[0]，切片 word[0:5]",
        'name = "Kerlaier"\nword = "Hello Python"\nword[0]  → "H"\nword[-1] → "n"\nword[0:5]  → "Hello"\nword[::-1] → "nohtyP olleH"',
        'word[0] = "H"\nword[0:5] = "Hello"\nword[::-1] = "nohtyP olleH"',
        "字符串是字符的序列，索引从0开始，切片超级灵活")
    add_code_block(doc,
        's = "  Hello, Python World!  "\n'
        's.strip()           # "Hello, Python World!"  去首尾空格\n'
        's.upper()           # "  HELLO, PYTHON WORLD!  "  全大写\n'
        's.lower()           # "  hello, python world!  "  全小写\n'
        's.replace("World", "朋友")  # 替换\n'
        '"apple,banana".split(",")  # ["apple", "banana"]  分割\n'
        'len(s)              # 24  字符串长度\n'
        '\n# f-string（重要！）\n'
        'name = "Kerlaier"\n'
        'print(f"我叫{name}，今年{25}岁")',
        "字符串常用方法 + f-string")
    add_tip_card(doc, "笔记", "f-string 是最常用的格式化方式：f\"我叫{name}\"，比 + 拼接直观十倍")

    # 三、布尔
    add_section_title(doc, "三、布尔类型：bool")
    add_knowledge_card(doc,
        "True / False（首字母大写！）",
        "is_raining = True\nprint(10 > 3)   # True\nprint(10 == 3)  # False",
        "10 > 3 → True\n10 == 3 → False",
        "布尔值只有 True 和 False，比较运算的结果就是布尔值")
    add_tip_card(doc, "避坑", "一个等号 = 是赋值，两个等号 == 是比较！初学者最容易犯的错")
    add_code_block(doc,
        "# 逻辑运算\n"
        "True and False  → False   # 全真才真\n"
        "True or False   → True    # 有真就真\n"
        "not True        → False   # 取反\n\n"
        "# bool() 判断\n"
        "bool(0) → False  bool(1) → True\n"
        "bool(\"\") → False  bool(\"Hi\") → True",
        "逻辑运算 & bool()")

    # 四、类型转换
    add_section_title(doc, "四、类型转换")
    add_knowledge_card(doc,
        "int(x) / float(x) / str(x) / bool(x)",
        'num_str = "42"\nint(num_str)    # 42\nfloat("3.14")   # 3.14\nstr(100)        # "100"',
        'int("42") → 42\nfloat("3.14") → 3.14\nstr(100) → "100"',
        "input() 返回的永远是字符串，要算数必须先 int() 转换")
    add_tip_card(doc, "避坑", '"5" + "3" = "53"（字符串拼接！），5 + 3 = 8（数字加法）。想算数，先 int()！')

    # ======== 作业部分 ========
    add_section_title(doc, "✏️ 课后练习")

    answers = parse_practice_answers(2)
    for num, question, code, hint in answers:
        add_homework_item(doc, num, question, code, hint)



    add_checklist(doc, [
        "能区分 int 和 float，会用基本算术运算符",
        "能对字符串做索引、切片、常用方法操作",
        "会用 f-string 格式化输出",
        "理解 True/False 和比较运算、逻辑运算",
        "掌握 int()/str()/float() 类型转换",
        "完成全部 5 道练习题 ✅"
    ])

    add_watermark(doc)
    os.makedirs(OUTPUT, exist_ok=True)
    path = os.path.join(OUTPUT, "第2课_数据类型.docx")
    doc.save(path)
    print(f"OK: {path}")
    copy_source_files(2)
    return path

# ============================================================
# 第3课：输入输出 & 运算符
# ============================================================
def generate_lesson_3():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21); section.page_height = Cm(29.7)
    section.left_margin = Cm(2); section.right_margin = Cm(2)
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2)

    add_chapter_banner(doc, 3, "输入输出 & 运算符", "2026-07-15")

    p = doc.add_paragraph()
    run = p.add_run("🎯 本章目标")
    run.bold = True; run.font.size = Pt(13); run.font.color.rgb = THEME_BLUE
    for g in ["掌握 print() 输出和 input() 输入", "理解 input() 永远返回字符串这条黄金规则",
              "学会用 int()/float() 转换后做运算", "实战项目：餐厅点餐计算器"]:
        p = doc.add_paragraph(f"  ✓ {g}")
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)

    add_section_title(doc, "一、print() —— 输出内容到屏幕")
    add_knowledge_card(doc,
        "print(内容1, 内容2, sep=\"分隔符\", end=\"结尾符\")",
        'print("菜品:", "红烧肉", "| 价格:", 38, "元")\nprint("正在计算", end="...")\nprint("完成！")',
        '菜品: 红烧肉 | 价格: 38 元\n正在计算...完成！',
        "print() 是最常用的输出函数，sep 控制分隔符，end 控制结尾")

    add_code_block(doc,
        'dish = "红烧肉"\nprice = 38\n'
        'print("菜品:", dish, "| 价格:", price, "元")\n\n'
        '# sep 控制分隔符\n'
        'print("姓名", "张三", "金额", "100", sep=" | ")\n'
        '# 输出: 姓名 | 张三 | 金额 | 100\n\n'
        '# end 控制结尾（默认换行）\n'
        'print("正在计算", end="...")\nprint("完成！")\n'
        '# 输出: 正在计算...完成！',
        "print() 实战用法")

    add_section_title(doc, "二、input() —— 从键盘获取输入")
    add_knowledge_card(doc,
        "变量 = input(\"提示文字\")\n⚠️ input() 返回的永远是字符串！",
        'name = input("请输入你的名字: ")\nprint(f"欢迎，{name}！")',
        '请输入你的名字: 小明\n欢迎，小明！',
        "input() 让程序能和用户交互，但返回值是 str，需要 int() 转换才能算数")

    add_tip_card(doc, "避坑", 'age = input("年龄: ")  # 你输入 25\n'
               'print(age + 1)   # 报错！因为 "25" + 1 不能算\n'
               'age = int(input("年龄: "))  # 正确：先转成数字再算')

    add_section_title(doc, "三、类型转换 —— 把 input 的字符串变成数字")
    add_knowledge_card(doc,
        "int(字符串) / float(字符串) —— 字符串转数字",
        'price_str = "38"\ncount_str = "2"\nprice = int(price_str)   # 38\ncount = int(count_str)   # 2\ntotal = price * count      # 76',
        '单价: 38元\n数量: 2\n小计: 76元',
        "input() → str → int() → 运算，这是处理用户输入的标准流程")

    add_section_title(doc, "四、运算符速查")
    add_code_block(doc,
        "x, y = 10, 3\n\n"
        "# 算术运算符\n"
        "x + y   → 13    # 加\n"
        "x - y   → 7     # 减\n"
        "x * y   → 30    # 乘\n"
        "x / y   → 3.33  # 除（有小数）\n"
        "x // y  → 3     # 整除（不要小数）\n"
        "x % y   → 1     # 取余（剩下的）\n"
        "x ** y  → 1000  # 次方\n\n"
        "# 比较运算符（结果是 True/False）\n"
        "x == y  → False  # 等于（双等号！）\n"
        "x != y  → True   # 不等于\n"
        "x > y   → True   # 大于\n\n"
        "# 赋值运算符\n"
        "x += 5  等价于 x = x + 5\n"
        "x -= 5  等价于 x = x - 5",
        "运算符速查表")

    add_section_title(doc, "五、实战：餐厅点餐计算器")
    add_code_block(doc,
        "# 完整账单计算\n"
        "dish_price = 38        # 一道菜的价格\n"
        "dish_count = 2         # 点了两份\n"
        "drink_price = 5        # 饮料单价\n"
        "drink_count = 3        # 点了三杯\n"
        "service_fee_rate = 0.1  # 服务费 10%\n\n"
        "food_total = dish_price * dish_count\n"
        "drink_total = drink_price * drink_count\n"
        "subtotal = food_total + drink_total\n"
        "service_fee = subtotal * service_fee_rate\n"
        "grand_total = subtotal + service_fee\n\n"
        "print(f\"菜品: {dish_price}元 x {dish_count} = {food_total}元\")\n"
        "print(f\"饮料: {drink_price}元 x {drink_count} = {drink_total}元\")\n"
        "print(f\"小计: {subtotal}元\")\n"
        "print(f\"服务费({service_fee_rate*100}%): {service_fee:.1f}元\")\n"
        "print(f\"应付总额: {grand_total:.1f}元\")\n\n"
        "# AA制分账\n"
        "people = 3\n"
        "per_person = grand_total // people   # 整除\n"
        "remainder = grand_total % people     # 取余\n"
        "print(f\"{people}人AA: 每人付{per_person:.0f}元, 多出{remainder:.1f}元\")",
        "完整点餐计算器")

    add_section_title(doc, "✏️ 课后练习")
    answers = parse_practice_answers(3)
    for num, question, code, hint in answers:
        add_homework_item(doc, num, question, code, hint)



    add_checklist(doc, [
        "会用 print() 输出，理解 sep 和 end 参数",
        "会用 input() 获取用户输入",
        "牢记：input() 返回的永远是字符串",
        "掌握 int()/float() 转换后做运算",
        "熟练掌握算术和比较运算符",
        "独立完成餐厅账单计算器"
    ])

    add_watermark(doc)
    os.makedirs(OUTPUT, exist_ok=True)
    path = os.path.join(OUTPUT, "第3课_输入输出.docx")
    doc.save(path)
    print(f"OK: {path}")
    copy_source_files(3)

# ============================================================
# 第4课：条件判断 if/elif/else
# ============================================================
def generate_lesson_4():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21); section.page_height = Cm(29.7)
    section.left_margin = Cm(2); section.right_margin = Cm(2)
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2)

    add_chapter_banner(doc, 4, "条件判断", "2026-07-16")

    p = doc.add_paragraph()
    run = p.add_run("🎯 本章目标")
    run.bold = True; run.font.size = Pt(13); run.font.color.rgb = THEME_BLUE
    for g in ["掌握 if / elif / else 多分支判断", "学会用比较运算和逻辑运算组合条件",
              "理解缩进决定代码归属", "实战：会员折扣系统、登录验证、成绩评级"]:
        p = doc.add_paragraph(f"  ✓ {g}")
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)

    add_section_title(doc, "一、if —— 满足条件才执行")
    add_knowledge_card(doc,
        "if 条件:\n    缩进4格，条件为 True 时执行的代码",
        "score = 75\nif score >= 60:\n    print(\"及格了！\")  # 75>=60，会执行",
        "及格了！",
        "if 后面跟条件，条件成立就执行缩进的代码块")

    add_section_title(doc, "二、if/else —— 二选一")
    add_knowledge_card(doc,
        "if 条件:\n    条件为 True 时执行\nelse:\n    条件为 False 时执行",
        "age = 16\nif age >= 18:\n    print(\"已成年\")\nelse:\n    print(\"未成年\")",
        "未成年",
        "两条路必然走一条，else 是 if 的兜底方案")

    add_section_title(doc, "三、if/elif/else —— 多选一")
    add_knowledge_card(doc,
        "if 条件1:\n    走这里\nelif 条件2:\n    走这里\nelif 条件3:\n    走这里\nelse:\n    都不满足走这里",
        "score = 85\nif score >= 90: grade = \"A\"\nelif score >= 80: grade = \"B\"\nelif score >= 70: grade = \"C\"\nelse: grade = \"D\"",
        "分数 85 分，等级：B",
        "从上到下依次判断，命中一条就跳出，后面的不再看")

    add_tip_card(doc, "避坑", "即使 85 也满足 \">=70\"，但程序只走第一个命中的 \">=80\"。所以条件的顺序很重要！")

    add_section_title(doc, "四、实战：商场会员折扣系统")
    add_code_block(doc,
        "# 会员折扣规则：\n"
        "#   钻石会员 → 7折\n"
        "#   金卡会员 → 8折\n"
        "#   银卡会员 → 9折\n"
        "#   普通用户 → 不打折\n"
        "#   消费满500额外减50\n\n"
        "member_level = \"金卡\"\n"
        "amount = 600\n\n"
        "if member_level == \"钻石\":\n"
        "    rate = 0.7\n"
        "elif member_level == \"金卡\":\n"
        "    rate = 0.8\n"
        "elif member_level == \"银卡\":\n"
        "    rate = 0.9\n"
        "else:\n"
        "    rate = 1.0\n\n"
        "discounted = amount * rate\n"
        "if discounted >= 500:\n"
        "    discounted -= 50\n\n"
        "print(f\"最终付款：{discounted}元\")",
        "会员折扣系统")

    add_section_title(doc, "五、条件组合 —— and / or / not")
    add_knowledge_card(doc,
        "and: 两个条件都满足\nor: 任一条件满足\nnot: 取反",
        'username = "admin"\npassword = "123456"\nif username == "admin" and password == "123456":\n    print("登录成功")\nelse:\n    print("用户名或密码错误")',
        "登录成功",
        "and = 且，or = 或，not = 非，用来组合多个条件")

    add_code_block(doc,
        "# 判断是否是 VIP 或消费超过1000\n"
        "is_vip = False\n"
        "amount = 1200\n"
        "if is_vip or amount >= 1000:\n"
        '    print("享受优先配送！")\n'
        "# 不满足VIP，但消费够了 → 也能享受",
        "or 实战：满足一个就行")

    add_section_title(doc, "✏️ 课后练习")
    answers = parse_practice_answers(4)
    for num, question, code, hint in answers:
        add_homework_item(doc, num, question, code, hint)



    add_checklist(doc, [
        "理解 if / elif / else 的执行逻辑",
        "知道从上到下依次判断，命中一条就跳出",
        "会用 and / or / not 组合条件",
        "理解缩进决定代码是否属于 if",
        "能独立写会员折扣、登录验证、成绩评级",
        "完成全部课后练习 ✅"
    ])

    add_watermark(doc)
    os.makedirs(OUTPUT, exist_ok=True)
    path = os.path.join(OUTPUT, "第4课_条件判断.docx")
    doc.save(path)
    print(f"OK: {path}")
    copy_source_files(4)

# ============================================================
# 第5课：循环 for / while
# ============================================================
def generate_lesson_5():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21); section.page_height = Cm(29.7)
    section.left_margin = Cm(2); section.right_margin = Cm(2)
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2)

    add_chapter_banner(doc, 5, "循环", "2026-07-16")

    p = doc.add_paragraph()
    run = p.add_run("🎯 本章目标")
    run.bold = True; run.font.size = Pt(13); run.font.color.rgb = THEME_BLUE
    for g in ["掌握 for 循环遍历集合", "掌握 while 循环按条件重复",
              "会用 break 提前退出和 continue 跳过", "实战：猜数字游戏、密码重试、乘法表"]:
        p = doc.add_paragraph(f"  ✓ {g}")
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)

    add_section_title(doc, "一、for 循环 —— 对集合里每个东西都做一遍")
    add_knowledge_card(doc,
        "for 变量 in 集合:\n    重复执行的代码（缩进）",
        'cart = ["苹果", "香蕉", "牛奶", "面包"]\nfor item in cart:\n    print(f"✓ {item}")',
        '✓ 苹果\n✓ 香蕉\n✓ 牛奶\n✓ 面包',
        "for 就像购物清单——把每样东西挨个念一遍")

    add_code_block(doc,
        "# range(起始, 结束, 步长) —— 重复固定次数\n"
        "# range(1, 10) → 1, 2, 3, ..., 9（不包含10）\n"
        "for i in range(1, 10):\n"
        "    print(f\"3 x {i} = {3 * i}\")\n\n"
        "# 输出：\n"
        "# 3 x 1 = 3\n"
        "# 3 x 2 = 6\n"
        "# ...\n"
        "# 3 x 9 = 27",
        "range() 打印乘法表")

    add_section_title(doc, "二、while 循环 —— 条件满足就一直做")
    add_knowledge_card(doc,
        "while 条件:\n    条件为 True 就一直重复",
        "count = 5\nwhile count > 0:\n    print(f\"倒计时：{count}\")\n    count = count - 1\nprint(\"发射！\")",
        '倒计时：5\n倒计时：4\n倒计时：3\n倒计时：2\n倒计时：1\n发射！',
        "while 是不确定次数的循环——等条件变成 False 才停")

    add_tip_card(doc, "避坑", "while 循环一定要有退出条件！忘记 count = count - 1 就会无限循环，程序卡死")

    add_section_title(doc, "三、for vs while 一句话区分")
    add_code_block(doc,
        "# for  = 知道要做多少次 → 数着做\n"
        "#       例：打印列表里每个物品\n"
        "# while = 不知道多少次，等条件变False → 等着做\n"
        "#       例：一直猜数字直到猜对\n\n"
        '# 口诀：\n'
        '#   "遍历"用 for，"等待"用 while',
        "for vs while 选择指南")

    add_section_title(doc, "四、break —— 提前退出循环")
    add_knowledge_card(doc,
        "break —— 立即跳出整个循环，后面的不做了",
        'shelf = ["雪碧", "芬达", "可乐", "矿泉水"]\nfor drink in shelf:\n    print(f"查看：{drink}")\n    if drink == "可乐":\n        print("找到了！")\n        break',
        '查看：雪碧\n查看：芬达\n查看：可乐\n找到了！',
        'break 就像：找到了就不找了，直接跳出循环')

    add_section_title(doc, "五、continue —— 跳过当前，继续下一个")
    add_knowledge_card(doc,
        "continue —— 跳过本轮循环，继续下一轮",
        'products = [("牛奶", 15), ("大米", 180), ("鸡蛋", 7)]\nfor name, days in products:\n    if days < 30:\n        continue\n    print(f"{name}：合格")',
        '大米：合格',
        'continue 就像：这个不行，下一个 —— 跳过不满足条件的')

    add_section_title(doc, "六、实战：猜数字游戏")
    add_code_block(doc,
        'import random\n'
        'target = random.randint(1, 100)  # 随机生成1-100\n'
        'attempts = 0\n\n'
        'while True:\n'
        '    guess = int(input("猜一个数字(1-100): "))\n'
        '    attempts += 1\n'
        '    if guess > target:\n'
        '        print("太大了！")\n'
        '    elif guess < target:\n'
        '        print("太小了！")\n'
        '    else:\n'
        '        print(f"猜对了！你用{attempts}次猜中了")\n'
        '        break',
        "猜数字游戏（while + break + if）")

    add_section_title(doc, "七、实战：密码重试（最多3次）")
    add_code_block(doc,
        'correct_password = "123456"\n'
        'max_try = 3\n'
        'tried = 0\n\n'
        'while tried < max_try:\n'
        '    pw = input("请输入密码：")\n'
        '    tried += 1\n'
        '    if pw == correct_password:\n'
        '        print(f"登录成功！（第{tried}次尝试）")\n'
        '        break\n'
        '    else:\n'
        f'        remaining = max_try - tried\n'
        f'        if remaining > 0:\n'
        '            print(f"密码错误，还剩{remaining}次机会")\n'
        f'        else:\n'
        '            print("3次都用完了，账号已锁定！")',
        "密码重试（while + break + 计数器）")

    add_section_title(doc, "✏️ 课后练习")
    answers = parse_practice_answers(5)
    for num, question, code, hint in answers:
        add_homework_item(doc, num, question, code, hint)



    add_checklist(doc, [
        "理解 for 循环：遍历集合或 range() 固定次数",
        "理解 while 循环：条件为 True 就一直重复",
        "掌握 break（提前退出）和 continue（跳过当前）",
        "能区分什么时候用 for，什么时候用 while",
        "能独立写猜数字游戏和密码重试逻辑",
        "完成全部课后练习"
    ])

    add_watermark(doc)
    os.makedirs(OUTPUT, exist_ok=True)
    path = os.path.join(OUTPUT, "第5课_循环.docx")
    doc.save(path)
    print(f"OK: {path}")
    copy_source_files(5)


# ============================================================
# 第6课：列表 list
# ============================================================
def generate_lesson_6():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21); section.page_height = Cm(29.7)

    add_chapter_banner(doc, 6, "列表 list", "2026-07-17")

    add_section_title(doc, "一、什么是列表？——能装多个东西的容器")
    add_code_block(doc,
        "# 语法：用 [] 创建，里面用逗号隔开\n"
        'cart = ["苹果", "香蕉", "牛奶"]    # 字符串列表\n'
        "scores = [95, 88, 73, 90]          # 数字列表\n"
        'mixed = ["小明", 20, True, 95.5]   # 混合类型（可以不常用）\n\n'
        'print(f"购物车: {cart}")\n'
        'print(f"成绩单: {scores}")\n'
        'print(f"有几样东西: {len(cart)}件")  # len() 查元素个数',
        "场景1：购物车")

    add_knowledge_card(doc,
        "列表 = []  里面放多个值，用逗号隔开",
        'fruits = ["苹果", "香蕉", "橘子"]\nscores = [95, 88, 73, 90]',
        '["苹果", "香蕉", "橘子"]\n[95, 88, 73, 90]',
        "列表就是一个能装多个东西的容器，用 [] 包起来")

    add_section_title(doc, "二、索引和切片 —— 和字符串一模一样")
    add_knowledge_card(doc,
        "列表[索引]   索引从0开始，-1是最后一个\n列表[起点:终点]   切片不包含终点",
        'scores = [95, 88, 73, 90]\nprint(scores[0])      # 第一个：95\nprint(scores[-1])     # 最后一个：90\nprint(scores[1:3])    # 索引1到2: [88, 73]\nprint(scores[:2])     # 前两个: [95, 88]',
        '95\n90\n[88, 73]\n[95, 88]',
        "列表的索引和切片规则跟字符串完全一样")

    add_section_title(doc, "三、增加元素 —— append / insert")
    add_knowledge_card(doc,
        "列表.append(值)    加到末尾\n列表.insert(位置, 值)    插到指定位置",
        'cart = ["苹果", "香蕉"]\ncart.append("牛奶")       # 加到最后\nprint(cart)               # ["苹果","香蕉","牛奶"]\ncart.insert(1, "草莓")    # 插到1号位\nprint(cart)               # ["苹果","草莓","香蕉","牛奶"]',
        '["苹果", "香蕉", "牛奶"]\n["苹果", "草莓", "香蕉", "牛奶"]',
        "append 像排队站最后，insert 像插队到指定位置")

    add_section_title(doc, "四、删除元素 —— remove / pop")
    add_knowledge_card(doc,
        "列表.remove(值)    按内容删除\n列表.pop()    弹掉最后一个（返回它）\n列表.pop(索引)    弹掉指定位置",
        'cart = ["苹果", "香蕉", "牛奶", "面包"]\ncart.remove("香蕉")      # 按内容删\nprint(cart)               # ["苹果","牛奶","面包"]\npopped = cart.pop()       # 弹走最后一个\nprint(f"弹出: {popped}")  # 面包\ncart.pop(0)               # 弹走第0个\nprint(cart)               # ["牛奶"]',
        '["苹果", "牛奶", "面包"]\n弹出: 面包\n["牛奶"]',
        "remove 是点名删除，pop 是弹走并拿到手")

    add_tip_card(doc, "避坑", "remove() 删的是第一个匹配的值，如果有重复只删第一个。pop() 返回被删除的值，可以接着用。")

    add_section_title(doc, "五、修改元素 —— 直接赋值")
    add_knowledge_card(doc,
        "列表[索引] = 新值",
        'scores = [95, 88, 73, 90]\nscores[2] = 80           # 把第3个改成80\nprint(scores)              # [95, 88, 80, 90]',
        '[95, 88, 80, 90]',
        "通过索引直接赋值就能修改，跟变量赋值一样简单")

    add_section_title(doc, "六、常用操作 —— 排序、查找、遍历")
    add_knowledge_card(doc,
        "sort() 排序 | reverse() 反转 | in 判断存在 | for 遍历",
        'scores = [95, 88, 73, 90, 85]\nscores.sort()            # 从小到大\nprint(scores)              # [73,85,88,90,95]\nscores.sort(reverse=True)  # 从大到小\nprint(90 in scores)        # True\nfor s in scores:\n    print(s)',
        '[73, 85, 88, 90, 95]\nTrue\n逐行打印每个分数',
        "sort 排序、in 判断存在、for 遍历是列表三大日常操作")

    add_section_title(doc, "七、实战：待办清单管理器")
    add_code_block(doc,
        'todos = ["写作业", "跑步"]                 # 初始清单\nprint(f"待办: {todos}")\n\n'
        'todos.append("买菜")                       # 新增事项\nprint(f"增加后: {todos}")\n\n'
        'todos.remove("跑步")                       # 完成"跑步"，删掉\nprint(f"完成后: {todos}")\n\n'
        'print(f"\n共{len(todos)}件事待办:")\n'
        'for i, task in enumerate(todos, 1):        # enumerate 同时拿索引和值\n'
        '    print(f"  {i}. {task}")               # 编号从1开始',
        "待办清单（append + remove + enumerate）")

    add_tip_card(doc, "笔记", "enumerate(列表, 1) 可以在遍历时同时拿到序号和内容，做待办清单特别好用。")

    add_section_title(doc, "八、列表操作速查表")
    add_code_block(doc,
        "创建：    lst = [1, 2, 3]\n"
        "查个数：  len(lst)              # 3\n"
        "加最后：  lst.append(4)         # [1, 2, 3, 4]\n"
        "插中间：  lst.insert(1, \"a\")    # [1, 'a', 2, 3]\n"
        "按值删：  lst.remove(2)         # [1, 3]\n"
        "弹最后：  x = lst.pop()         # 取出最后一个\n"
        "弹位置：  x = lst.pop(0)        # 取出第1个\n"
        "修改：    lst[0] = \"新值\"\n"
        "排序：    lst.sort()            # 从小到大\n"
        "反转：    lst.reverse()\n"
        "找位置：  lst.index(3)          # 3在第几个位置\n"
        "出现次数：lst.count(3)\n"
        "是否存在：3 in lst              # True/False\n"
        "遍历：    for item in lst:",
        "列表操作速查表")

    add_section_title(doc, "✏️ 课后练习")
    answers = parse_practice_answers(6)
    for num, question, code, hint in answers:
        add_homework_item(doc, num, question, code, hint)

    add_checklist(doc, [
        "理解列表是什么：用 [] 装多个数据的容器",
        "掌握索引 [0]、[-1] 和切片 [1:3]",
        "掌握 append() 追加和 insert() 插入",
        "掌握 remove() 按值删除和 pop() 弹出",
        "会用 sort() 排序和 in 判断存在",
        "会 for 循环遍历列表",
        "能独立写待办清单管理器",
        "完成全部课后练习"
    ])

    add_watermark(doc)
    os.makedirs(OUTPUT, exist_ok=True)
    path = os.path.join(OUTPUT, "第6课_列表.docx")
    doc.save(path)
    print(f"OK: {path}")
    copy_source_files(6)


if __name__ == "__main__":
    generate_lesson_2()
    generate_lesson_3()
    generate_lesson_4()
    generate_lesson_5()
    generate_lesson_6()
    print("\n全部生成完毕！")


